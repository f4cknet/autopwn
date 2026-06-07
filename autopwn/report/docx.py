"""AutoPwn report layer — DOCX report generator.

Moved from ``autopwn/_legacy.py`` v3.1 in P3.2 (see ``rebuild.md`` §4.4 +
§6.4).  Reads from the typed :class:`ExploitInfo` dataclass (P3.1) instead
of the loose ``exploit_info`` dict, and writes to ``out_dir`` instead of
the current working directory.

P3.6 update (2026-06-07)
========================
``python-docx`` is now imported **lazily at module level with
``try/except ImportError``** (was function-level in P3.2).  If the
library is missing, ``generate_docx`` falls back to a markdown
generator (``_generate_markdown``) that writes ``{target}_wp.md`` to
``out_dir`` instead of the .docx.  See ``refactor.md`` §10.

Why module-level, not function-level
------------------------------------
The ``rebuild.md`` §6.4 spec example has the ``try/except ImportError``
at the *caller* (record_success) level.  That works for function-level
imports (re-raise at call time), but module-level imports fail at
*import time* — so the caller's ``except ImportError`` would never
trigger (the import error happens when ``autopwn.report.docx`` is
imported, before ``record_success`` is ever called).

The fix: do the ``try/except ImportError`` at module top, set a
``_HAS_DOCX = True/False`` module flag, and have ``generate_docx``
dispatch to the markdown fallback when ``_HAS_DOCX`` is False.  This
matches the spec's *intent* (markdown fallback when docx missing)
without the latent import-time crash that would otherwise break the
whole ``autopwn.report`` package.

Design
======
* **Module-level import guard**: ``_HAS_DOCX`` flag is ``True`` if
  ``python-docx`` imports cleanly, ``False`` otherwise.
* **Pure-ish function**: takes ``info`` and ``out_dir``, returns the
  generated ``Path``.  Caller (P3.4 ``record_success``) is responsible
  for success-gating (``--no-report``).
* **No more global state**: the legacy ``global exploit_info`` declaration
  is gone.  All 14 ``exploit_info['x']`` reads become ``info.x``.

Field mapping (legacy dict → ExploitInfo)
-----------------------------------------
========================================== ==============================
``exploit_info['x']``                       ``info.x`` (or derived)
========================================== ==============================
``'target_binary'``                         ``info.target_binary``
                                           (basename extracted via
                                           ``Path().stem`` in this fn)
``'exploit_type'``                          ``info.exploit_type``
``'architecture'``                          ``info.architecture``
``'vulnerability_type'``                    ``info.vulnerability_type``
``'padding'``                               ``info.padding``
``'addresses'``                             ``info.addresses``
``'payload'`` (bytes → hex; str → str)      ``info.payload``
``'timestamp'``                             ``info.timestamp``
``'success'``                               *removed* (ExploitInfo is
                                           constructed only on success;
                                           gate is P3.4 record_success)
========================================== ==============================

Adoption roadmap (see ``rebuild.md`` §4.4)
------------------------------------------
* P3.1 (✅) — :class:`ExploitInfo` dataclass.
* P3.2 (✅) — move ``generate_docx_report`` here as
  ``generate_docx(info, out_dir)``.
* P3.3 (✅) — code generator moved to ``autopwn/report/code.py``.
* P3.6 (this PR) — ``try/except ImportError`` markdown fallback.
* P3.4 — refactor ``handle_exploitation_success`` to construct an
  ``ExploitInfo`` directly and call ``report.record_success`` (no
  more dict bridge).
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Optional

from autopwn.core.logging import Colors, print_error, print_success, VERSION
from autopwn.report.model import ExploitInfo
from autopwn.report.code import generate_code

# P3.6: lazy / safe python-docx import.  If the library is missing,
# fall back to a markdown report.  See module docstring for the
# rationale on module-level (vs function-level) import.
try:
    from docx import Document  # noqa: F401
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401
    _HAS_DOCX = True
except ImportError:
    Document = None  # type: ignore[assignment]
    WD_ALIGN_PARAGRAPH = None  # type: ignore[assignment]
    _HAS_DOCX = False


def _generate_markdown(info: ExploitInfo, out_dir: Path) -> Path:
    """Write a markdown report as a fallback when ``python-docx`` is missing.

    The output filename mirrors the docx layout (``{target}_wp.md``)
    so the caller can find the report regardless of which generator
    ran.  The content covers the same 5 sections as the docx
    (Basic Information, Buffer Overflow Information, Key Address
    Information, Exploitation Code, Exploitation Summary) plus a
    footer.
    """
    target_name = Path(info.target_binary).name
    if target_name.startswith("./"):
        target_name = target_name[2:]
    target_name = Path(target_name).stem
    report_path = out_dir / f"{target_name}_wp.md"

    lines: list[str] = []
    lines.append(f"# PWN Exploitation Report — {info.target_binary}\n")
    lines.append("## Basic Information\n")
    lines.append(f"- **Target Binary**: {info.target_binary}")
    lines.append(f"- **Exploitation Time**: {info.timestamp}")
    lines.append(f"- **Architecture**: {info.architecture}")
    lines.append(f"- **Vulnerability Type**: {info.vulnerability_type}")
    lines.append(f"- **Exploitation Method**: {info.exploit_type}\n")

    lines.append("## Buffer Overflow Information\n")
    lines.append(f"- **Buffer Overflow Padding**: {info.padding} bytes\n")

    if info.addresses:
        lines.append("## Key Address Information\n")
        lines.append("| Address Type | Address Value |")
        lines.append("| --- | --- |")
        for addr_type, addr_value in info.addresses.items():
            if isinstance(addr_value, int):
                v = f"0x{addr_value:x}"
            elif isinstance(addr_value, str) and addr_value.isdigit():
                v = f"0x{int(addr_value):x}"
            elif isinstance(addr_value, str) and addr_value.startswith("0x"):
                v = addr_value
            else:
                try:
                    v = f"0x{int(str(addr_value)):x}" if "x" not in str(addr_value) else str(addr_value)
                except Exception:
                    v = str(addr_value)
            lines.append(f"| {addr_type} | {v} |")
        lines.append("")

    if info.payload:
        lines.append("## Exploitation Code\n")
        lines.append("```python")
        # P3.3: code generator; out_dir may be used in P3.4 / P3.5 to
        # write a .py file artifact.
        lines.append(generate_code(info, out_dir).rstrip())
        lines.append("```\n")
        if isinstance(info.payload, bytes):
            lines.append(f"**Payload Length**: {len(info.payload)} bytes\n")
        else:
            lines.append(f"**Payload Length**: {len(str(info.payload))} characters\n")

    lines.append("## Exploitation Summary\n")
    lines.append("- **Exploitation Status**: Successful")
    lines.append(
        f"- **Exploitation Method**: Successfully gained shell access through "
        f"{info.vulnerability_type} vulnerability using {info.exploit_type} technique.\n"
    )

    lines.append("---\n")
    lines.append(f"_Report Generation Tool_: AutoPwn v{VERSION}  ")
    lines.append(f"_Generation Time_: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("")
    lines.append("> **Note**: Markdown fallback — `python-docx` is not installed.  "
                 "Install with `pip install python-docx` for a richer .docx report.")

    report_path.write_text("\n".join(lines), encoding="utf-8")
    return report_path


def generate_docx(info: ExploitInfo, out_dir: Path) -> Optional[Path]:
    """Generate a DOCX (or markdown fallback) exploitation report.

    Parameters
    ----------
    info : ExploitInfo
        Typed result of a successful exploitation run.  All 8 fields
        are read; ``success`` gate is **not** checked here (caller
        responsibility — P3.4 ``record_success``).
    out_dir : Path
        Directory to write ``{target}_wp.docx`` (or ``.md`` if
        ``python-docx`` is missing) into.

    Returns
    -------
    Path or None
        Path of the generated report on success; ``None`` on any
        failure (matches the legacy behavior of swallowing exceptions
        and printing an error).

    Side effects
    ------------
    * Writes ``{target}_wp.docx`` (or ``.md``) to ``out_dir``.
    * Prints success / error messages via ``core.logging``.
    """
    # P3.6: markdown fallback dispatch.  If python-docx is missing,
    # delegate to _generate_markdown and return early.
    if not _HAS_DOCX:
        try:
            report_path = _generate_markdown(info, out_dir)
            print_success(
                f"Exploitation report generated (markdown fallback): "
                f"{Colors.YELLOW}{report_path}{Colors.END}"
            )
            return report_path
        except Exception as e:
            print_error(f"Failed to generate markdown report: {e}")
            return None

    try:
        # Extract target name: basename without path / extension.
        # Legacy behavior strips a leading "./" defensively.
        target_name = Path(info.target_binary).name
        if target_name.startswith("./"):
            target_name = target_name[2:]
        target_name = Path(target_name).stem  # strip extension
        report_filename = f"{target_name}_wp.docx"
        report_path = out_dir / report_filename

        # Build the document
        doc = Document()

        # Title
        title = doc.add_heading("PWN Exploitation Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Basic information
        doc.add_heading("Basic Information", level=1)
        basic_info = doc.add_paragraph()
        basic_info.add_run("Target Binary: ").bold = True
        basic_info.add_run(f"{info.target_binary}\n")
        basic_info.add_run("Exploitation Time: ").bold = True
        basic_info.add_run(f"{info.timestamp}\n")
        basic_info.add_run("Architecture: ").bold = True
        basic_info.add_run(f"{info.architecture}\n")
        basic_info.add_run("Vulnerability Type: ").bold = True
        basic_info.add_run(f"{info.vulnerability_type}\n")
        basic_info.add_run("Exploitation Method: ").bold = True
        basic_info.add_run(f"{info.exploit_type}\n")

        # Buffer overflow information
        doc.add_heading("Buffer Overflow Information", level=1)
        padding_info = doc.add_paragraph()
        padding_info.add_run("Buffer Overflow Padding: ").bold = True
        padding_info.add_run(f"{info.padding} bytes\n")

        # Key address information (table)
        if info.addresses:
            doc.add_heading("Key Address Information", level=1)
            addr_table = doc.add_table(rows=1, cols=2)
            addr_table.style = "Table Grid"
            hdr_cells = addr_table.rows[0].cells
            hdr_cells[0].text = "Address Type"
            hdr_cells[1].text = "Address Value"

            for addr_type, addr_value in info.addresses.items():
                row_cells = addr_table.add_row().cells
                row_cells[0].text = addr_type
                # Hex formatting (verbatim from legacy L288-302)
                if isinstance(addr_value, int):
                    row_cells[1].text = f"0x{addr_value:x}"
                elif isinstance(addr_value, str) and addr_value.isdigit():
                    row_cells[1].text = f"0x{int(addr_value):x}"
                elif isinstance(addr_value, str) and addr_value.startswith("0x"):
                    row_cells[1].text = addr_value
                else:
                    try:
                        if "x" in str(addr_value):
                            row_cells[1].text = str(addr_value)
                        else:
                            row_cells[1].text = f"0x{int(str(addr_value)):x}"
                    except Exception:
                        row_cells[1].text = str(addr_value)

        # Exploitation code (only if we have a payload)
        if info.payload:
            doc.add_heading("Exploitation Code", level=1)
            payload_para = doc.add_paragraph()
            payload_para.add_run("Complete Python Exploitation Code:\n").bold = True

            # P3.3: code generator moved to report.code; signature is
            # (info, out_dir) -> str (forward-compat: out_dir may be
            # used in P3.4 / P3.5 to write a .py file artifact).
            exploitation_code = generate_code(info, out_dir)
            payload_para.add_run(f"{exploitation_code}\n")

            payload_para.add_run("Payload Length: ").bold = True
            if isinstance(info.payload, bytes):
                payload_para.add_run(f"{len(info.payload)} bytes\n")
            else:
                payload_para.add_run(f"{len(str(info.payload))} characters\n")

        # Exploitation summary
        doc.add_heading("Exploitation Summary", level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run("Exploitation Status: ").bold = True
        summary_para.add_run("Successful\n")
        summary_para.add_run("Exploitation Method: ").bold = True
        summary_para.add_run(
            f"Successfully gained shell access through {info.vulnerability_type} "
            f"vulnerability using {info.exploit_type} technique.\n"
        )

        # Footer
        doc.add_paragraph("\n" + "─" * 50)
        footer_para = doc.add_paragraph()
        footer_para.add_run("Report Generation Tool: ").bold = True
        footer_para.add_run(f"AutoPwn v{VERSION}\n")
        footer_para.add_run("Generation Time: ").bold = True
        footer_para.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # Save
        doc.save(str(report_path))
        print_success(
            f"Exploitation report generated: {Colors.YELLOW}{report_path}{Colors.END}"
        )
        return report_path

    except Exception as e:
        print_error(f"Failed to generate report: {e}")
        return None


__all__ = ["generate_docx", "_generate_markdown"]
